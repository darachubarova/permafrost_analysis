#!/usr/bin/env python3
"""
Convert Diploma/ГЛАВА_2_КОРР.md → Diploma/ГЛАВА_2_КОРР.docx
with GOST-compliant typography for Word.

Steps:
1. Generate pandoc default reference.docx
2. Patch it (Times New Roman 14pt, 1.5 line spacing, GOST margins)
3. Run pandoc with the patched reference.docx
4. Post-process: fix margins on every section in the output file
"""

import subprocess
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PANDOC      = "/tmp/pandoc-3.6.4/bin/pandoc"
WORKSPACE   = "/workspaces/permafrost_analysis_dara"
MD_FILE     = os.path.join(WORKSPACE, "Diploma", "ГЛАВА_2_КОРР.md")
DOCX_FILE   = os.path.join(WORKSPACE, "Diploma", "ГЛАВА_2_КОРР.docx")
REF_DOCX    = "/tmp/gost_reference.docx"

# ── 1. Generate default reference.docx ──────────────────────────────────────
print("Generating default reference.docx …")
with open(REF_DOCX, "wb") as f:
    subprocess.run([PANDOC, "--print-default-data-file", "reference.docx"],
                   stdout=f, check=True)

# ── 2. Patch the reference with python-docx ──────────────────────────────────
print("Patching reference styles …")
ref = Document(REF_DOCX)

# GOST margins: left 30 mm, right 15 mm, top 20 mm, bottom 20 mm
for section in ref.sections:
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Normal / Body Text style
FONT_BODY = "Times New Roman"
SIZE_BODY = Pt(14)

def patch_style(style, font_name, font_size, bold=False, align=None, space_before=None, space_after=None):
    style.font.name = font_name
    style.font.size = font_size
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = Pt(21)   # ~1.5 × 14 pt
    pf.first_line_indent = Cm(1.25)   # GOST paragraph indent
    if align:
        pf.alignment = align
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)

style_names = [s.name for s in ref.styles]

if "Normal" in style_names:
    patch_style(ref.styles["Normal"], FONT_BODY, SIZE_BODY,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY)

if "Body Text" in style_names:
    patch_style(ref.styles["Body Text"], FONT_BODY, SIZE_BODY,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Heading styles (1–4)
heading_sizes = {1: Pt(16), 2: Pt(15), 3: Pt(14), 4: Pt(14)}
for lvl, size in heading_sizes.items():
    name = f"Heading {lvl}"
    if name in style_names:
        s = ref.styles[name]
        s.font.name = FONT_BODY
        s.font.size = size
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        pf = s.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if lvl <= 2 else WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(12)
        pf.space_after  = Pt(6)
        pf.line_spacing = Pt(21)
        pf.first_line_indent = Cm(0)

# Table style: no special patch needed – pandoc uses Table style
ref.save(REF_DOCX)
print(f"  Saved patched reference → {REF_DOCX}")

# ── 3. Run pandoc ─────────────────────────────────────────────────────────────
print("Running pandoc conversion …")
cmd = [
    PANDOC,
    MD_FILE,
    "-o", DOCX_FILE,
    "--reference-doc", REF_DOCX,
    "--from", "markdown+pipe_tables+raw_html+auto_identifiers+tex_math_dollars",
    "--to", "docx",
    "--resource-path", os.path.join(WORKSPACE, "Diploma")
                     + ":" + os.path.join(WORKSPACE, "figures")
                     + ":" + WORKSPACE,
    "--wrap", "none",
    "--standalone",
    "--toc",
    "--toc-depth", "3",
]
subprocess.run(cmd, check=True, cwd=os.path.join(WORKSPACE, "Diploma"))
print(f"  Created: {DOCX_FILE}")

# ── 4. Post-process output docx (fix margins, page numbers) ──────────────────
print("Post-processing output docx …")
doc = Document(DOCX_FILE)

for section in doc.sections:
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Ensure all paragraph fonts are Times New Roman (pandoc sometimes leaves runs un-styled)
for para in doc.paragraphs:
    for run in para.runs:
        if not run.font.name:
            run.font.name = FONT_BODY

doc.save(DOCX_FILE)
print("  Post-processing done.")

print(f"\nAll done → {DOCX_FILE}")
